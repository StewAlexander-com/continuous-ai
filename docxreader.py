"""docxreader — deterministic Word (.docx) text extraction for :read.

Uses python-docx to extract paragraphs, headings, and tables from Office Open
XML packages. Layout, images, and embedded objects are not invented — missing
text is refused plainly. Legacy ``.doc`` (Word 97–2003) is out of scope;
convert to ``.docx`` or PDF first.
"""
from __future__ import annotations

import os
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

DEFAULT_DOCX_MAX_PARAS = 20_000
DEFAULT_DOCX_MAX_TABLES = 500
MIN_EXTRACTABLE_CHARS = 20


@dataclass(frozen=True)
class DocxOptions:
    enabled: bool = True
    max_paragraphs: int = DEFAULT_DOCX_MAX_PARAS
    max_tables: int = DEFAULT_DOCX_MAX_TABLES


def docx_options_from_config(config: dict[str, Any] | None) -> DocxOptions:
    c = config or {}
    return DocxOptions(
        enabled=bool(c.get("docx_reader_enabled", True)),
        max_paragraphs=int(c.get("docx_max_paragraphs", DEFAULT_DOCX_MAX_PARAS)),
        max_tables=int(c.get("docx_max_tables", DEFAULT_DOCX_MAX_TABLES)),
    )


def _import_docx():
    try:
        import docx  # python-docx
        return docx
    except ImportError:
        return None


def docx_runtime_ready() -> bool:
    return _import_docx() is not None


def _looks_like_docx_zip(path: Path) -> bool:
    """True when the file is a ZIP that claims to be an OOXML package."""
    try:
        if not zipfile.is_zipfile(path):
            return False
        with zipfile.ZipFile(path, "r") as zf:
            names = set(zf.namelist())
        return "[Content_Types].xml" in names and any(
            n.startswith("word/") for n in names
        )
    except (OSError, zipfile.BadZipFile):
        return False


def _paragraph_lines(doc, *, cap: int) -> tuple[list[str], int, bool]:
    """Collect non-empty paragraph texts up to ``cap``. Returns (lines, total, truncated)."""
    lines: list[str] = []
    total = 0
    truncated = False
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        total += 1
        if len(lines) >= cap:
            truncated = True
            continue
        style = ""
        try:
            style_name = (para.style.name or "").strip() if para.style else ""
            if style_name.lower().startswith("heading"):
                style = f"[{style_name}] "
        except Exception:
            style = ""
        lines.append(f"{style}{text}")
    return lines, total, truncated


def _table_blocks(doc, *, cap: int) -> tuple[list[str], int, bool]:
    """Format tables as TSV-like blocks. Returns (blocks, total, truncated)."""
    blocks: list[str] = []
    total = 0
    truncated = False
    for table in doc.tables:
        total += 1
        if len(blocks) >= cap:
            truncated = True
            continue
        rows: list[str] = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(
                    (p.text or "").strip() for p in cell.paragraphs
                ).strip()
                cells.append(cell_text.replace("\t", " ").replace("\n", " "))
            # Skip fully empty rows.
            if any(cells):
                rows.append("\t".join(cells))
        if not rows:
            continue
        blocks.append(
            f"--- Table {len(blocks) + 1} ({len(rows)} row(s)) ---\n"
            + "\n".join(rows)
        )
    return blocks, total, truncated


def extract_docx(path: Path, *, opts: DocxOptions | None = None) -> tuple[bool, str, str]:
    """Extract text from a user-named .docx. Returns (ok, name_or_err, text)."""
    options = opts or DocxOptions()
    name = path.name
    if not options.enabled:
        return False, ("DOCX reading is disabled in config.yaml "
                       "(set docx_reader_enabled: true)."), ""

    docx_mod = _import_docx()
    if docx_mod is None:
        return False, ("DOCX support needs python-docx: "
                       "pip install python-docx  (or re-run bash setup.sh)"), ""

    if not _looks_like_docx_zip(path):
        return False, (f"{name} is not a readable .docx package "
                       "(corrupt, renamed, or not Office Open XML). "
                       "I won't guess its contents."), ""

    try:
        doc = docx_mod.Document(str(path))
    except Exception as e:
        err = str(e).lower()
        if "password" in err or "encrypt" in err:
            return False, (f"{name} appears password-protected — I can't read "
                           "encrypted Word files. Decrypt it or export to PDF/text."), ""
        return False, f"Cannot open {name} as a .docx: {e}", ""

    para_cap = max(1, options.max_paragraphs)
    table_cap = max(0, options.max_tables)
    para_lines, para_total, para_trunc = _paragraph_lines(doc, cap=para_cap)
    table_blocks, table_total, table_trunc = _table_blocks(doc, cap=table_cap)

    body_parts: list[str] = []
    if para_lines:
        body_parts.append("--- Body ---\n" + "\n\n".join(para_lines))
    if table_blocks:
        body_parts.append("\n\n".join(table_blocks))
    body = "\n\n".join(body_parts).strip()
    total_chars = len(body)

    if total_chars < MIN_EXTRACTABLE_CHARS:
        return False, (f"{name} has no extractable text "
                       f"({para_total} paragraph(s), {table_total} table(s) checked; "
                       "likely empty or image-only). I won't guess its contents."), ""

    core = getattr(doc, "core_properties", None)
    title = ""
    author = ""
    if core is not None:
        title = (getattr(core, "title", None) or "").strip()
        author = (getattr(core, "author", None) or "").strip()

    header_lines = [
        "[DOCX DOCUMENT PROFILE]",
        f"File: {name}",
        f"Paragraphs: {para_total} total"
        + (f" (showing first {len(para_lines)})" if para_trunc else ""),
        f"Tables: {table_total} total"
        + (f" (showing first {len(table_blocks)})" if table_trunc else ""),
        f"Characters extracted: {total_chars:,}",
        "Method: python-docx paragraph + table extraction",
    ]
    if title:
        header_lines.append(f"Title: {title}")
    if author:
        header_lines.append(f"Author: {author}")
    header_lines.append(
        "[NOTICE: Layout, images, charts, and embedded objects are not included. "
        "Do not invent content that was not extracted as text.]"
    )
    if para_trunc or table_trunc:
        header_lines.append(
            "[CONTENT TRUNCATION: extraction caps were hit — do not claim knowledge "
            "of omitted paragraphs/tables. Raise docx_max_paragraphs / "
            "docx_max_tables in config.yaml, or attach an excerpt.]"
        )

    return True, name, "\n".join(header_lines) + "\n\n" + body


def load_docx(path_str: str, max_mb: int | None = None,
              opts: DocxOptions | None = None) -> tuple[bool, str, str]:
    """Validate + extract a user-named .docx. Size check mirrors load_file()."""
    _DEFAULT_MAX_MB = 50
    if not path_str or not path_str.strip():
        return False, "No file path given. Usage: :read <path>", ""
    p = Path(os.path.expanduser(path_str.strip()))
    if not p.exists():
        return False, f"No file at {p} -- check the path.", ""
    if not p.is_file():
        return False, f"{p} is not a regular file.", ""
    try:
        size = p.stat().st_size
    except OSError as e:
        return False, f"Cannot stat {p}: {e}", ""
    limit = int((max_mb if max_mb is not None else _DEFAULT_MAX_MB) * 1024 * 1024)
    if size > limit:
        return False, (f"{p.name} is {size/1024/1024:.1f} MB -- over the {limit//1024//1024} MB "
                       "attach limit. Raise max_attach_mb in config.yaml, or attach an excerpt."), ""
    return extract_docx(p, opts=opts)


def legacy_doc_refusal(path_str: str) -> str:
    """Honest message for Word 97–2003 ``.doc`` (not OOXML)."""
    name = Path(path_str).name or path_str
    return (f"{name} is a legacy .doc (Word 97–2003) file — I only extract "
            ".docx (Office Open XML). Save/export as .docx or PDF, then :read again.")


def format_attachment_readers_status_lines(config: dict[str, Any] | None = None) -> list[str]:
    """Printable lines for :setup / :status — PDF + DOCX readiness."""
    c = config or {}
    lines = ["  ── Attachments (:read) ──"]

    pdf_on = bool(c.get("pdf_reader_enabled", True))
    try:
        from pdfreader import _import_fitz, _ocr_runtime_ready
        pdf_ready = _import_fitz() is not None
        ocr_ready = _ocr_runtime_ready()
    except Exception:
        pdf_ready = False
        ocr_ready = False
    if not pdf_on:
        lines.append("  PDF:        disabled (pdf_reader_enabled: false)")
    elif pdf_ready:
        ocr = "OCR ready" if ocr_ready else "OCR optional (install tesseract + pytesseract)"
        lines.append(f"  PDF:        OK — PyMuPDF ({ocr})")
    else:
        lines.append("  PDF:        NEEDS FIX — pip install pymupdf  (or bash setup.sh)")

    docx_on = bool(c.get("docx_reader_enabled", True))
    if not docx_on:
        lines.append("  DOCX:       disabled (docx_reader_enabled: false)")
    elif docx_runtime_ready():
        lines.append("  DOCX:       OK — python-docx (.docx only; not legacy .doc)")
    else:
        lines.append("  DOCX:       NEEDS FIX — pip install python-docx  (or bash setup.sh)")

    lines.append("  Tip:        :read ~/path/file.docx   |   legacy .doc → save as .docx/PDF")
    return lines
