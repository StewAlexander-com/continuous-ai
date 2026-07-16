"""Tests for filereader — honest reading of user-attached files (:read).

Covers the honesty-critical behaviors: real contents are returned, truncation is
ALWAYS announced (so the model can't characterize unseen content), binary/missing
files are refused plainly (never guessed), and CSV is summarized structurally.
Pure/deterministic — no model, no network.
"""
import os
import sys
import tempfile

import filereader as fr


def _tmp(content: bytes, suffix=".txt"):
    fd, path = tempfile.mkstemp(suffix=suffix)
    os.write(fd, content)
    os.close(fd)
    return path


def test_reads_real_text_contents():
    p = _tmp(b"hello\nworld\n", ".txt")
    try:
        ok, block = fr.read_attachment(p)
        assert ok
        assert "hello" in block and "world" in block
        assert "USER-ATTACHED FILE" in block
        assert "Citation contract" in block
        assert "hypothesize" in block.lower() or "beyond the attachment" in block.lower()
        assert "Answer only from" not in block
    finally:
        os.unlink(p)
    print("[PASS] reads real text contents and marks them as user-attached")


def test_python_file_reads():
    p = _tmp(b"def foo():\n    return 42\n", ".py")
    try:
        ok, block = fr.read_attachment(p)
        assert ok and "def foo" in block
    finally:
        os.unlink(p)
    print("[PASS] reads .py files")


def test_truncation_is_announced():
    # A file larger than the default budget, read via read_attachment (chunk 1),
    # must carry an explicit paging notice in its BODY (not just the header).
    big = ("\n".join(f"line {i}" for i in range(20000))).encode()
    p = _tmp(big, ".txt")
    try:
        ok, block = fr.read_attachment(p)   # uses DEFAULT_BUDGET_CHARS
        assert ok
        assert "PAGING / TRUNCATION NOTICE" in block, "long file must carry a paging notice"
        assert "not been shown" in block.lower()
        assert ":more" in block
    finally:
        os.unlink(p)
    print("[PASS] a long file's first chunk carries an explicit paging notice")


def test_missing_file_refused_plainly():
    ok, msg = fr.read_attachment("/nonexistent/path/to/file.txt")
    assert not ok and "No file" in msg
    print("[PASS] missing file refused with an honest message (no guessed contents)")


def test_binary_refused():
    p = _tmp(b"\x00\x01\x02\x03BINARY\x00\xff", ".bin")
    try:
        ok, msg = fr.read_attachment(p)
        assert not ok
        assert "binary" in msg.lower() or "cannot read" in msg.lower()
        assert "won't guess" in msg.lower()
        assert fr.should_offer_read_miss_menu(p) is False
    finally:
        os.unlink(p)
    print("[PASS] binary file refused, contents never guessed")


def test_miss_menu_gated_for_existing_unreadable():
    """Existing binary/oversize must not claim 'No file' or self-suggest."""
    import shutil
    d = tempfile.mkdtemp()
    try:
        binary = os.path.join(d, "notes.bear2bk")
        with open(binary, "wb") as f:
            f.write(b"PK\x03\x04" + b"\x00" * 64)
        ok, msg, _ = fr.load_path(binary)
        assert not ok
        assert "No file or directory" not in msg
        assert "binary" in msg.lower() or "cannot read" in msg.lower()
        assert fr.should_offer_read_miss_menu(binary) is False
        # Defense: rank must not re-offer the exact existing path.
        cands = fr.rank_path_candidates(binary)
        assert not any(fr.paths_same_target(c, binary) for c in cands)
        missing = os.path.join(d, "notes-missing.bear2bk")
        assert fr.should_offer_read_miss_menu(missing) is True
    finally:
        shutil.rmtree(d)
    print("[PASS] miss menu gated; existing binary not self-suggested")


def test_permission_denied_honest_message():
    import shutil
    d = tempfile.mkdtemp()
    p = os.path.join(d, "secret.txt")
    with open(p, "w") as f:
        f.write("nope\n")
    try:
        os.chmod(p, 0)
        ok, msg, _ = fr.load_file(p)
        # Some environments (root/CI) may still read mode 0 — skip then.
        if ok:
            print("[SKIP] permission denied — process can read mode-0 file")
            return
        assert "permission" in msg.lower()
        assert fr.should_offer_read_miss_menu(p) is False
    finally:
        os.chmod(p, 0o644)
        shutil.rmtree(d)
    print("[PASS] permission denial uses honest message, no miss menu")


def test_drop_existing_pick_candidate_breaks_loop():
    import shutil
    d = tempfile.mkdtemp()
    try:
        a = os.path.join(d, "a.txt")
        b = os.path.join(d, "b.bin")
        open(a, "w").write("hi")
        open(b, "wb").write(b"\x00\xff")
        left = fr.drop_existing_pick_candidate([a, b], b)
        assert left == [a]
        gone = os.path.join(d, "gone.txt")
        left2 = fr.drop_existing_pick_candidate([a, gone], gone)
        assert gone in left2 and a in left2
    finally:
        shutil.rmtree(d)
    print("[PASS] drop_existing_pick_candidate removes only still-present picks")


def test_empty_path_refused():
    ok, msg = fr.read_attachment("")
    assert not ok and "Usage" in msg
    print("[PASS] empty path gives usage, not a crash")


def test_csv_small_shows_all():
    csv_bytes = b"name,age\nAlice,30\nBob,25\n"
    p = _tmp(csv_bytes, ".csv")
    try:
        ok, block = fr.read_attachment(p)
        assert ok
        assert "2 data row" in block and "Alice" in block and "Bob" in block
        assert "name (text)" in block and "age (int)" in block
        # The instructional header mentions the word 'TRUNCATION' generically;
        # what must be ABSENT for a small file is the actual notice marker.
        assert "TRUNCATION NOTICE" not in block, "small CSV should not be truncated"
    finally:
        os.unlink(p)
    print("[PASS] small CSV: full table + inferred column types, no truncation")


def test_csv_large_sampled_with_notice():
    rows = "name,n\n" + "\n".join(f"r{i},{i}" for i in range(fr.CSV_FULL_ROWS + 100))
    p = _tmp(rows.encode(), ".csv")
    try:
        ok, block = fr.read_attachment(p)
        assert ok
        assert "TRUNCATION NOTICE" in block, "large CSV must announce sampling"
        assert f"{fr.CSV_FULL_ROWS + 100} data row" in block
        assert "Do not state totals" in block
    finally:
        os.unlink(p)
    print("[PASS] large CSV: sampled with an explicit 'do not state totals' notice")


def test_oversize_refused():
    # Use a tiny custom max_mb so we don't have to write 50 MB.
    big = b"x" * (2 * 1024 * 1024)  # 2 MB
    p = _tmp(big, ".txt")
    try:
        ok, msg, _ = fr.load_file(p, max_mb=1)   # 1 MB limit
        assert not ok and "over the" in msg.lower() and "limit" in msg.lower()
    finally:
        os.unlink(p)
    print("[PASS] oversize file refused before reading (configurable limit)")


def test_default_accept_limit_is_50mb():
    assert fr.DEFAULT_MAX_ATTACH_MB == 50
    assert fr.max_attach_bytes() == 50 * 1024 * 1024
    assert fr.max_attach_bytes(10) == 10 * 1024 * 1024
    print("[PASS] default attach limit is 50 MB, configurable")


def test_budget_scales_with_num_ctx():
    small = fr.budget_chars(None)          # default when unknown
    big = fr.budget_chars(32768)           # large context
    assert big > small, "budget should grow with num_ctx"
    assert fr.budget_chars(0) == fr.DEFAULT_BUDGET_CHARS
    assert fr.budget_chars(100) >= fr.MIN_BUDGET_CHARS, "honest floor on tiny context"
    print("[PASS] per-chunk budget scales with num_ctx, with an honest floor")


def test_read_chunk_paging_advances_and_completes():
    text = "\n".join(f"line {i}" for i in range(5000))
    budget = 2000
    seen = 0
    offset = 0
    chunks = 0
    done = False
    while not done and chunks < 1000:
        c = fr.read_chunk(text, "big.txt", char_offset=offset, budget=budget)
        assert c["next_offset"] > offset or c["done"], "offset must advance"
        seen += c["shown_chars"]
        offset = c["next_offset"]
        done = c["done"]
        chunks += 1
    assert done, "paging must eventually complete"
    assert seen >= len(text) - 5, "paging must cover ~all chars (no silent loss)"
    assert chunks > 1, "a 5000-line file should take multiple chunks at budget=2000"
    print("[PASS] read_chunk pages forward, covers the whole file, then reports done")


def test_partial_chunk_announces_paging():
    text = "\n".join(f"line {i}" for i in range(5000))
    c = fr.read_chunk(text, "big.txt", char_offset=0, budget=2000)
    assert not c["done"]
    assert "PAGING" in c["block"] and ":more" in c["block"]
    assert "NOT been shown" in c["block"]
    print("[PASS] a partial chunk explicitly announces unseen content + :more")


def test_small_file_one_chunk_no_notice():
    c = fr.read_chunk("just a little text\n", "tiny.txt", char_offset=0, budget=8000)
    assert c["done"]
    # The instructional header mentions 'PAGING/TRUNCATION' generically; what must
    # be absent for a one-chunk file is the actual notice marker + the FINAL CHUNK
    # banner (a single complete chunk needs neither).
    assert "NOTICE \u2014" not in c["block"] and "FINAL CHUNK" not in c["block"]
    assert ":more" not in c["block"]
    print("[PASS] a small file fits in one chunk with no paging notice")


def test_suggest_lists_nearby_files():
    import tempfile, pathlib
    d = tempfile.mkdtemp()
    open(os.path.join(d, "voice.py"), "w").close()
    open(os.path.join(d, "session.py"), "w").close()
    try:
        # ask for a missing file in that dir -> error should suggest neighbors
        ok, msg, _ = fr.load_file(os.path.join(d, "voce.py"))  # typo
        assert not ok
        assert "Nearby" in msg and "voice.py" in msg
    finally:
        import shutil; shutil.rmtree(d)
    print("[PASS] not-found error suggests nearby files (did-you-mean)")


def test_rank_path_candidates_fuzzy_typo():
    import shutil
    d = tempfile.mkdtemp()
    voice = os.path.join(d, "voice.py")
    open(voice, "w").close()
    open(os.path.join(d, "session.py"), "w").close()
    try:
        cands = fr.rank_path_candidates(os.path.join(d, "voce.py"))
        assert cands and voice in cands[0] or any("voice.py" in c for c in cands)
        assert "session.py" not in cands[0] or len(cands) > 1
    finally:
        shutil.rmtree(d)
    print("[PASS] rank_path_candidates surfaces typo-close filenames")


def test_rank_path_candidates_glob_miss():
    import shutil
    d = tempfile.mkdtemp()
    open(os.path.join(d, "alpha.py"), "w").close()
    open(os.path.join(d, "beta.txt"), "w").close()
    try:
        cands = fr.rank_path_candidates(os.path.join(d, "*.py"))
        assert len(cands) == 1 and cands[0].endswith("alpha.py")
    finally:
        shutil.rmtree(d)
    print("[PASS] rank_path_candidates lists glob extension matches")


def test_parse_read_pick_single_yes():
    cands = ["/tmp/voice.py"]
    assert fr.parse_read_pick_response("y", cands) == ("pick", cands[0])
    assert fr.parse_read_pick_response("1", cands) == ("pick", cands[0])
    assert fr.parse_read_pick_response("n", cands) == ("cancel", None)
    assert fr.parse_read_pick_response("hello", cands) == ("not_pick", None)
    print("[PASS] parse_read_pick_response handles single-candidate y/1/n")


def test_parse_read_pick_multi_requires_number():
    cands = ["/tmp/a.py", "/tmp/b.py"]
    assert fr.parse_read_pick_response("y", cands) == ("retry", None)
    assert fr.parse_read_pick_response("2", cands) == ("pick", cands[1])
    assert fr.parse_read_pick_response("99", cands) == ("retry", None)
    assert fr.parse_read_pick_response("", cands) == ("not_pick", None)
    assert fr.parse_read_pick_response("b.py", cands) == ("pick", cands[1])
    menu = fr.format_read_pick_menu("/tmp/voce.py", cands)
    assert "Press Return" in menu
    print("[PASS] multi-candidate pick rejects bare y, accepts valid index")


def test_parse_read_pick_poka_yoke_retry_and_basename():
    cands = ["/proj/index.html", "/proj/assets/"]
    labels = ["index.html", "assets/"]
    assert fr.parse_read_pick_response(
        "0", cands, mode="directory", labels=labels) == ("retry", None)
    assert fr.parse_read_pick_response(
        "3", cands, mode="directory", labels=labels) == ("retry", None)
    assert fr.parse_read_pick_response(
        "index.html", cands, mode="directory", labels=labels) == ("pick", cands[0])
    assert fr.parse_read_pick_response(
        "assets/", cands, mode="directory", labels=labels) == ("pick", cands[1])
    print("[PASS] out-of-range retries; unique names pick")


def test_list_directory_and_load_path():
    import tempfile
    d = tempfile.mkdtemp()
    open(os.path.join(d, "alpha.txt"), "w").close()
    os.mkdir(os.path.join(d, "subdir"))
    try:
        ok, name, text = fr.list_directory(d)
        assert ok and "alpha.txt" in text and "subdir/" in text
        ok2, name2, text2 = fr.load_path(d)
        assert ok2 and text2 == text and "(directory listing)" in name2
        ok3, name3, text3 = fr.load_path(os.path.join(d, "alpha.txt"))
        assert ok3 and name3 == "alpha.txt"
    finally:
        import shutil; shutil.rmtree(d)
    print("[PASS] load_path lists directories and reads files")


def test_soft_absolute_users_missing_slash():
    target = "/Users/stewartalexander/stewalexander-com-git/network-attack-visualizer/index.html"
    if not os.path.isfile(target):
        print("[SKIP] soft absolute — sample path not on this machine")
        return
    missing = target.lstrip("/")
    resolved, notice = fr.resolve_read_path(missing)
    assert resolved == target, resolved
    assert notice and "added leading /" in notice
    # Exact relative that does NOT exist under cwd must still soft-correct.
    ok, name, text = fr.load_path(missing)
    assert ok and name == "index.html" and "html" in text.lower()[:200]
    # Existing relative path must win over soft absolute (non-regression).
    import tempfile, shutil
    d = tempfile.mkdtemp()
    try:
        os.makedirs(os.path.join(d, "Users", "demo"))
        f = os.path.join(d, "Users", "demo", "note.txt")
        with open(f, "w") as fh:
            fh.write("relative wins\n")
        old = os.getcwd()
        os.chdir(d)
        try:
            r, n = fr.resolve_read_path("Users/demo/note.txt")
            assert n is None and os.path.samefile(r, f)
        finally:
            os.chdir(old)
    finally:
        shutil.rmtree(d)
    print("[PASS] soft absolute Users/... → /Users/... (exact relative still wins)")


def test_suggest_index_stem_after_soft_absolute():
    parent = "/Users/stewartalexander/stewalexander-com-git/network-attack-visualizer"
    if not os.path.isdir(parent):
        print("[SKIP] index stem suggest — sample dir not on this machine")
        return
    attempted = f"Users/stewartalexander/stewalexander-com-git/network-attack-visualizer/index"
    cands = fr.rank_path_candidates(attempted)
    assert any(c.rstrip("/").endswith("index.html") for c in cands), cands
    print("[PASS] missing-slash + bare 'index' suggests index.html")


def test_directory_chunk_paging():
    # Force multi-chunk by exceeding MIN_BUDGET_CHARS with a long listing.
    lines = ["Directory listing for /tmp/demo (0 dir(s), 400 file(s) shown):", "", "Files:"]
    lines.extend(f"  file_with_a_reasonably_long_name_{i:04d}.txt" for i in range(400))
    text = "\n".join(lines)
    name = "demo/ (directory listing)"
    budget = fr.MIN_BUDGET_CHARS  # floor still applies inside read_chunk
    assert len(text) > budget
    c1 = fr.read_directory_chunk(text, name, char_offset=0, budget=budget)
    assert not c1["done"], "expected more than one chunk"
    assert "```" not in c1["block"]
    assert "PAGING" in c1["block"]
    offset = c1["next_offset"]
    done = False
    guard = 0
    while not done and guard < 50:
        cn = fr.read_directory_chunk(text, name, char_offset=offset, budget=budget)
        offset = cn["next_offset"]
        done = cn["done"]
        guard += 1
    assert done
    print("[PASS] directory listings page with :more (no code fence)")


def test_directory_browse_menu_and_return_review():
    import tempfile, shutil
    d = tempfile.mkdtemp()
    try:
        alpha = os.path.join(d, "alpha.txt")
        beta = os.path.join(d, "beta.md")
        subdir = os.path.join(d, "subdir")
        open(alpha, "w").close()
        open(beta, "w").close()
        os.mkdir(subdir)
        # Stable mtimes: newest must be menu item 1; oldest at the bottom.
        os.utime(alpha, (1_600_000_000, 1_600_000_000))
        os.utime(subdir, (1_650_000_000, 1_650_000_000))
        os.utime(beta, (1_700_000_000, 1_700_000_000))
        ok, dir_path, entries = fr.list_directory_entries(d)
        assert ok and len(entries) == 3
        labels = [e[0] for e in entries]
        assert "alpha.txt" in labels and "subdir/" in labels
        assert labels == ["beta.md", "subdir/", "alpha.txt"], labels
        menu = fr.format_directory_browse_menu(dir_path, entries)
        assert "Reply with a number (1–3)" in menu
        assert "Press Return to review the full directory listing" in menu
        assert "Out-of-range numbers keep this menu open" in menu
        assert "1  " in menu
        menu_lines = [line for line in menu.splitlines() if line.strip()[:1].isdigit()]
        assert menu_lines[0].endswith("beta.md")
        assert menu_lines[-1].endswith("alpha.txt")
        assert "2023-" in menu_lines[0] and "2020-" in menu_lines[-1]
        # Empty Return on directory mode → review (not dismiss)
        action, path = fr.parse_read_pick_response(
            "", [e[1] for e in entries], mode="directory")
        assert action == "review" and path is None
        # Number picks a path
        action2, path2 = fr.parse_read_pick_response(
            "1", [e[1] for e in entries], mode="directory")
        assert action2 == "pick" and path2 == entries[0][1]
        # Miss mode empty still dismisses
        action3, _ = fr.parse_read_pick_response("", ["x"], mode="miss")
        assert action3 == "not_pick"
    finally:
        shutil.rmtree(d)
    print("[PASS] directory browse menu + Return reviews listing")


def test_directory_mtime_portability_and_ties():
    import tempfile, shutil
    d = tempfile.mkdtemp()
    try:
        a = os.path.join(d, "Alpha.txt")
        z = os.path.join(d, "zeta.txt")
        folder = os.path.join(d, "Folder")
        open(a, "w").close()
        open(z, "w").close()
        os.mkdir(folder)
        same = 1_680_000_000
        os.utime(a, (same, same))
        os.utime(z, (same, same))
        os.utime(folder, (same, same))

        ok, _root, entries = fr.list_directory_entries(d)
        assert ok
        # Coarse/equal mtimes (common on Windows filesystems) use casefolded name.
        assert [label for label, _ in entries] == [
            "Alpha.txt", "Folder/", "zeta.txt"
        ]
        # Directory paths with the native trailing separator must stat correctly.
        folder_path = next(path for label, path in entries if label == "Folder/")
        assert folder_path.endswith(os.sep)
        assert fr.format_path_modified_time(folder_path) != "unknown modified"
        # Missing/unreadable timestamps are explicit, never crashes.
        assert fr.format_path_modified_time(
            os.path.join(d, "gone.txt")) == "unknown modified"
    finally:
        shutil.rmtree(d)
    print("[PASS] mtime sorting portable across coarse timestamps/native separators")


def test_directory_followup_resolves_named_direct_child_only():
    import tempfile, shutil
    d = tempfile.mkdtemp()
    outside = tempfile.mkdtemp()
    try:
        index = os.path.join(d, "index.html")
        readme = os.path.join(d, "README.md")
        with open(index, "w") as f:
            f.write("<h1>real content</h1>")
        with open(readme, "w") as f:
            f.write("# real readme")
        os.mkdir(os.path.join(d, "nested"))
        with open(os.path.join(d, "nested", "deep.js"), "w") as f:
            f.write("not a direct child")

        got = fr.resolve_directory_file_followup(
            "Review the index.html and summarize", d)
        assert got and os.path.samefile(got, index)

        # No action verb: mentioning a name alone must not trigger a read.
        assert fr.resolve_directory_file_followup(
            "The listing has index.html", d) is None
        # Multiple named files are ambiguous; do not silently choose one.
        assert fr.resolve_directory_file_followup(
            "Compare index.html and README.md", d) is None
        # No recursive inference.
        assert fr.resolve_directory_file_followup(
            "Review deep.js", d) is None

        # A symlink inside the named dir may point outside; inferred reads reject it.
        target = os.path.join(outside, "secret.txt")
        with open(target, "w") as f:
            f.write("outside")
        link = os.path.join(d, "shortcut.txt")
        try:
            os.symlink(target, link)
            assert fr.resolve_directory_file_followup(
                "Review shortcut.txt", d) is None
        except (OSError, NotImplementedError):
            pass
    finally:
        shutil.rmtree(d)
        shutil.rmtree(outside)
    print("[PASS] directory follow-up reattaches one explicit safe child")


def test_detect_local_read_intent():
    assert fr.detect_local_read_intent("can you read through what is at ~/?") == ("~", None)
    assert fr.detect_local_read_intent("read ~/foo.py") == ("~/foo.py", None)
    path, q = fr.detect_local_read_intent("read ~/foo.py and explain it")
    assert path == "~/foo.py" and q == "and explain it"
    assert fr.detect_local_read_intent("list ~/Documents") == ("~/Documents", None)
    assert fr.detect_local_read_intent("Read my GitHub at github.com/foo") is None
    assert fr.detect_local_read_intent("Summarize https://example.com/x") is None
    assert fr.detect_local_read_intent("hello there") is None
    print("[PASS] detect_local_read_intent matches local paths, rejects URLs")


def test_plain_read_path_with_spaces():
    import shutil
    base = tempfile.mkdtemp()
    spaced = os.path.join(base, "Misc Docs", "PDF Documents")
    os.makedirs(spaced)
    try:
        intent = fr.detect_local_read_intent(f"read {spaced}")
        assert intent is not None
        path, q = intent
        assert q is None
        assert os.path.samefile(os.path.expanduser(path), spaced)
        ok, name, text = fr.load_path(path)
        assert ok and "(directory listing)" in name
    finally:
        shutil.rmtree(base)
    print("[PASS] plain read resolves paths with spaces via longest-existing-prefix")


def test_plain_read_quoted_path_with_spaces():
    import shutil
    base = tempfile.mkdtemp()
    fpath = os.path.join(base, "My Notes.txt")
    with open(fpath, "w") as f:
        f.write("hello notes\n")
    try:
        path, q = fr.detect_local_read_intent(f'read "{fpath}" summarize')
        assert path == fpath and q == "summarize"
    finally:
        shutil.rmtree(base)
    print("[PASS] quoted plain-read paths with spaces still work")


def test_plain_read_glob_with_spaces_and_question():
    import shutil
    base = tempfile.mkdtemp()
    spaced = os.path.join(base, "Misc Docs", "PDF Documents")
    os.makedirs(spaced)
    for name in ("a.txt", "b.txt"):
        with open(os.path.join(spaced, name), "w") as f:
            f.write(f"notes {name}\n")
    try:
        glob_path = os.path.join(spaced, "*.txt")
        intent = fr.detect_local_read_intent(f"read {glob_path} any learned insights?")
        assert intent is not None
        path, q = intent
        assert path == glob_path
        assert q == "any learned insights?"
        ok, name, text = fr.load_path(path)
        assert ok and "2 file" in name
    finally:
        shutil.rmtree(base)
    print("[PASS] plain read glob with spaced path and trailing question")


def test_glob_expands_multiple_py_files():
    import shutil
    d = tempfile.mkdtemp()
    for name in ("alpha.py", "beta.py", "notes.txt"):
        with open(os.path.join(d, name), "w") as f:
            f.write(f"# {name}\n")
    try:
        ok, name, text = fr.load_path(os.path.join(d, "*.py"))
        assert ok, f"glob load failed: {name}"
        assert "2 file" in name
        assert "=== alpha.py ===" in text and "=== beta.py ===" in text
        assert "# notes.txt" not in text
        assert fr.expand_read_glob(os.path.join(d, "*.py")) != []
    finally:
        shutil.rmtree(d)
    print("[PASS] glob pattern expands to multiple .py files")


def test_glob_single_match_behaves_like_file():
    import shutil
    d = tempfile.mkdtemp()
    p = os.path.join(d, "only.py")
    with open(p, "w") as f:
        f.write("x = 1\n")
    try:
        ok, name, text = fr.load_path(os.path.join(d, "on*.py"))
        assert ok and name == "only.py" and "x = 1" in text
    finally:
        shutil.rmtree(d)
    print("[PASS] glob with one match delegates to single-file load")


def test_glob_no_match_honest_error():
    ok, msg, _ = fr.load_path("/nonexistent/dir/*.py")
    assert not ok and "No files match" in msg
    print("[PASS] unmatched glob returns honest error")


def test_literal_path_with_star_wins_over_glob():
    import shutil
    d = tempfile.mkdtemp()
    literal = os.path.join(d, "foo*bar.txt")
    with open(literal, "w") as f:
        f.write("literal star file\n")
    with open(os.path.join(d, "fooXbar.txt"), "w") as f:
        f.write("would match glob\n")
    try:
        ok, name, text = fr.load_path(literal)
        assert ok and "literal star" in text
    finally:
        shutil.rmtree(d)
    print("[PASS] existing literal path containing * is not glob-expanded")


def _make_pdf(path: str, pages: list[str]) -> None:
    import fitz
    doc = fitz.open()
    for text in pages:
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=11)
    doc.save(path)
    doc.close()


def test_pdf_extracts_native_text():
    import shutil
    try:
        import fitz  # noqa: F401
    except ImportError:
        print("[SKIP] pymupdf not installed")
        return
    d = tempfile.mkdtemp()
    pdf_path = os.path.join(d, "sample.pdf")
    try:
        _make_pdf(pdf_path, ["Hello PDF learning tool.", "Page two content here."])
        ok, name, text = fr.load_path(pdf_path)
        assert ok, f"PDF load failed: {name}"
        assert name == "sample.pdf"
        assert "PDF DOCUMENT PROFILE" in text
        assert "Hello PDF learning tool" in text
        assert "Page two content here" in text
        assert "--- Page 1 ---" in text and "--- Page 2 ---" in text
    finally:
        shutil.rmtree(d)
    print("[PASS] PDF native text extraction with page markers")


def test_pdf_encrypted_refused():
    import shutil
    try:
        import fitz
    except ImportError:
        print("[SKIP] pymupdf not installed")
        return
    d = tempfile.mkdtemp()
    enc_path = os.path.join(d, "secret.pdf")
    try:
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "hidden")
        doc.save(enc_path, encryption=fitz.PDF_ENCRYPT_AES_256, user_pw="nope")
        doc.close()
        ok, msg, _ = fr.load_path(enc_path)
        assert not ok and "password" in msg.lower()
    finally:
        shutil.rmtree(d)
    print("[PASS] encrypted PDF refused plainly")


def test_pdf_disabled_in_config():
    import shutil
    try:
        import fitz  # noqa: F401
    except ImportError:
        print("[SKIP] pymupdf not installed")
        return
    d = tempfile.mkdtemp()
    pdf_path = os.path.join(d, "x.pdf")
    try:
        _make_pdf(pdf_path, ["text"])
        ok, msg, _ = fr.load_file(pdf_path, pdf_options={"pdf_reader_enabled": False})
        assert not ok and "disabled" in msg.lower()
    finally:
        shutil.rmtree(d)
    print("[PASS] pdf_reader_enabled: false refuses PDF reads")


def test_pdf_pages_through_read_chunk():
    import shutil
    try:
        import fitz  # noqa: F401
    except ImportError:
        print("[SKIP] pymupdf not installed")
        return
    d = tempfile.mkdtemp()
    pdf_path = os.path.join(d, "big.pdf")
    try:
        pages = [
            "\n".join(f"page{p} line {i} padding text for paging" for i in range(400))
            for p in range(6)
        ]
        _make_pdf(pdf_path, pages)
        ok, name, text = fr.load_path(pdf_path)
        assert ok and len(text) > 8000, f"expected large extraction, got {len(text)}"
        chunk = fr.read_chunk(text, name, char_offset=0, budget=3000)
        assert not chunk["done"]
        assert "PAGING" in chunk["block"]
    finally:
        shutil.rmtree(d)
    print("[PASS] extracted PDF text pages via read_chunk / :more path")


def _make_docx(path: str, *, paragraphs: list[str] | None = None,
               table_rows: list[list[str]] | None = None) -> None:
    from docx import Document
    doc = Document()
    for text in paragraphs or []:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r_i, row in enumerate(table_rows):
            for c_i, cell in enumerate(row):
                table.rows[r_i].cells[c_i].text = cell
    doc.save(path)


def test_docx_extracts_paragraphs_and_tables():
    import shutil
    try:
        import docx  # noqa: F401
    except ImportError:
        print("[SKIP] python-docx not installed")
        return
    d = tempfile.mkdtemp()
    path = os.path.join(d, "sample.docx")
    try:
        _make_docx(
            path,
            paragraphs=["Emergency contact form.", "Primary: Jane Doe"],
            table_rows=[["Name", "Phone"], ["Jane", "555-0100"]],
        )
        ok, name, text = fr.load_path(path)
        assert ok, f"DOCX load failed: {name}"
        assert name == "sample.docx"
        assert "DOCX DOCUMENT PROFILE" in text
        assert "Emergency contact form" in text
        assert "Jane Doe" in text
        assert "555-0100" in text
        assert "--- Table 1" in text
        assert fr.should_offer_read_miss_menu(path) is False
    finally:
        shutil.rmtree(d)
    print("[PASS] DOCX paragraph + table extraction")


def test_docx_disabled_in_config():
    import shutil
    try:
        import docx  # noqa: F401
    except ImportError:
        print("[SKIP] python-docx not installed")
        return
    d = tempfile.mkdtemp()
    path = os.path.join(d, "off.docx")
    try:
        _make_docx(path, paragraphs=["Should not attach when disabled."])
        ok, msg, _ = fr.load_file(path, pdf_options={"docx_reader_enabled": False})
        assert not ok and "disabled" in msg.lower()
    finally:
        shutil.rmtree(d)
    print("[PASS] docx_reader_enabled: false refuses DOCX reads")


def test_docx_corrupt_refused():
    import shutil
    d = tempfile.mkdtemp()
    path = os.path.join(d, "fake.docx")
    try:
        with open(path, "wb") as f:
            f.write(b"PK\x03\x04not-a-real-docx")
        ok, msg, _ = fr.load_path(path)
        assert not ok
        assert "not a readable .docx" in msg.lower() or "cannot open" in msg.lower()
        assert "won't guess" in msg.lower() or "guess" in msg.lower()
    finally:
        shutil.rmtree(d)
    print("[PASS] corrupt/non-OOXML .docx refused plainly")


def test_legacy_doc_refused_with_convert_hint():
    import shutil
    d = tempfile.mkdtemp()
    path = os.path.join(d, "old.doc")
    try:
        with open(path, "wb") as f:
            f.write(b"\xd0\xcf\x11\xe0" + b"\x00" * 64)  # OLE-ish bytes
        ok, msg, _ = fr.load_path(path)
        assert not ok
        assert ".doc" in msg.lower() and "docx" in msg.lower()
        assert "save" in msg.lower() or "export" in msg.lower()
        assert fr.should_offer_read_miss_menu(path) is False
    finally:
        shutil.rmtree(d)
    print("[PASS] legacy .doc refused with convert-to-.docx hint")


def test_docx_pages_through_read_chunk():
    import shutil
    try:
        import docx  # noqa: F401
    except ImportError:
        print("[SKIP] python-docx not installed")
        return
    d = tempfile.mkdtemp()
    path = os.path.join(d, "big.docx")
    try:
        paras = [f"Paragraph {i} with padding text for paging checks." for i in range(400)]
        _make_docx(path, paragraphs=paras)
        ok, name, text = fr.load_path(path)
        assert ok and len(text) > 8000, f"expected large extraction, got {len(text)}"
        chunk = fr.read_chunk(text, name, char_offset=0, budget=2000)
        assert not chunk["done"]
        assert "PAGING" in chunk["block"]
    finally:
        shutil.rmtree(d)
    print("[PASS] extracted DOCX text pages via read_chunk / :more path")


def test_attachment_readers_status_mentions_docx():
    from docxreader import format_attachment_readers_status_lines
    lines = "\n".join(format_attachment_readers_status_lines({}))
    assert "DOCX" in lines and "PDF" in lines
    print("[PASS] attachment reader status lists PDF and DOCX")


def test_parse_read_arg_splits_path_and_question():
    # parser lives in seedling (CLI concern); import and exercise it.
    import seedling
    path, q = seedling._parse_read_arg("~/seedling/voice.py this is what gives you a voice")
    assert path == "~/seedling/voice.py", f"path mis-parsed: {path!r}"
    assert q == "this is what gives you a voice"
    # quoted path with spaces
    path2, q2 = seedling._parse_read_arg('"~/My Notes.txt" summarize it')
    assert path2 == "~/My Notes.txt" and q2 == "summarize it"
    # path only, no question
    path3, q3 = seedling._parse_read_arg("~/foo.py")
    assert path3 == "~/foo.py" and q3 is None
    path4, q4 = seedling._parse_read_arg(
        "/tmp/Misc Docs/PDF Documents/*.pdf any learned insights?")
    assert path4 == "/tmp/Misc Docs/PDF Documents/*.pdf"
    assert q4 == "any learned insights?"
    print("[PASS] :read arg splits into (path, question), quote-aware")


if __name__ == "__main__":
    tests = [v for k, v in sorted(globals().items())
             if k.startswith("test_") and callable(v)]
    failed = 0
    for t in tests:
        try:
            t()
        except AssertionError as e:
            failed += 1
            print(f"[FAIL] {t.__name__}: {e}")
        except Exception as e:
            failed += 1
            print(f"[ERROR] {t.__name__}: {type(e).__name__}: {e}")
    print(f"\n{len(tests) - failed}/{len(tests)} filereader checks passed")
    sys.exit(1 if failed else 0)
